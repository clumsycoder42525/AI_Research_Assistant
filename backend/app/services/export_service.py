from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document

class ExportService:
    @staticmethod
    def generate_pdf(report) -> BytesIO:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        # Title
        elements.append(Paragraph("Vedanta AI - Research Report", styles['Title']))
        elements.append(Spacer(1, 12))

        # Summary
        elements.append(Paragraph("Summary", styles['Heading2']))
        elements.append(Paragraph(report.summary, styles['Normal']))
        elements.append(Spacer(1, 12))

        # Key Points
        elements.append(Paragraph("Key Points", styles['Heading2']))
        elements.append(Paragraph(report.key_points.replace('\n', '<br/>'), styles['Normal']))
        elements.append(Spacer(1, 12))

        # Applications
        elements.append(Paragraph("Applications", styles['Heading2']))
        elements.append(Paragraph(report.applications, styles['Normal']))
        elements.append(Spacer(1, 12))

        # Conclusion
        elements.append(Paragraph("Conclusion", styles['Heading2']))
        elements.append(Paragraph(report.conclusion, styles['Normal']))
        elements.append(Spacer(1, 12))

        # Citations
        if report.citations:
            elements.append(Paragraph("References", styles['Heading2']))
            for cit in report.citations:
                elements.append(Paragraph(f"- {cit}", styles['Normal']))

        doc.build(elements)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_docx(report) -> BytesIO:
        doc = Document()
        doc.add_heading("Vedanta AI - Research Report", 0)

        doc.add_heading('Summary', level=1)
        doc.add_paragraph(report.summary)

        doc.add_heading('Key Points', level=1)
        doc.add_paragraph(report.key_points)

        doc.add_heading('Applications', level=1)
        doc.add_paragraph(report.applications)

        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph(report.conclusion)

        if report.citations:
            doc.add_heading('References', level=1)
            for cit in report.citations:
                doc.add_paragraph(cit, style='List Bullet')

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
